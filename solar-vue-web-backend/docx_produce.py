from pymongo import MongoClient
from bson.objectid import ObjectId
from datetime import datetime, date, timedelta, time
from dateutil.relativedelta import relativedelta
from io import BytesIO
from flask import send_file
from urllib.parse import quote

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_COLOR_INDEX
from docx.shared import RGBColor

from flask import current_app

import sys
import traceback
#-------------------------------------------------------------------------------
def exception_detail(e):
    error_class = e.__class__.__name__ #取得錯誤類型
    detail = e.args[0] #取得詳細內容
    cl, exc, tb = sys.exc_info() #取得Call Stack
    lastCallStack = traceback.extract_tb(tb)[-1] #取得Call Stack的最後一筆資料
    fileName = lastCallStack[0] #取得發生的檔案名稱
    lineNum = lastCallStack[1] #取得發生的行號
    funcName = lastCallStack[2] #取得發生的函數名稱
    errMsg = "File \"{}\", line {}, in {}: [{}] {}".format(fileName, lineNum, funcName, error_class, detail)
    return errMsg
#------------------------------------資料取得----------------------------------------
def mongo_connect():
    try:
        client = MongoClient("mongodb://root:pc152@61.64.52.109:27017/")
        return client["pv"]
    except Exception as e:
        print(e)

#輸入id，輸出plant_id, solar_ID, meter_ID, pr_ID
def id_identify(pv, id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    equipment = pv["equipment"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    if plant and not equipment:
        plant_id = id
        name = plant["name"]
        solar_ID = solar_meter_ID(pv, name)
        meter_ID = id
        pr_ID = id
        collection = plant.get("collection", " ")
    elif not plant and equipment:
        name = equipment["PV"]
        plant = pv["plant"].find_one(
            {
                "name": name
            }
        )
        collection = equipment.get("collection", " ")
        plant_id = str(plant["_id"])
        PV_str = equipment["PV"]
        if equipment["type"] == "pv_lgroup":
            pv_lgroup = equipment["name"]
            group = None
        elif equipment["type"] == "pv_group":
            pv_lgroup = equipment["lgroup"]
            group = equipment["name"]
        solar_ID = solar_meter_ID(pv, PV_str, pv_lgroup, group)
        meter_ID = id
        pr_ID = id
    return plant_id, solar_ID, meter_ID, pr_ID, collection

#案場資訊
def field_imformation(pv, object_id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(object_id)
        }
    )

    name = plant["name"]
    field_position = []
    field_position.append(plant["plant_address"])
    field_position.append(plant["coordinates"])
    capacity = plant["capacity"]
    return name, field_position, capacity

#日照計ID
def solar_meter_ID(pv, PV, lgroup=None, group=None):
    equipment = pv["equipment"]
    if lgroup == None and group == None:
        meter = equipment.find_one(
            {
                "PV": PV,
                "type": "sun",
                "main_sun": 1
            }
        )
    elif lgroup and not group:
        meter = equipment.find_one(
            {
                "PV": PV,
                "lgroup": lgroup,
                "type": "sun",
                "main_sun": 1
            }
        )
    elif lgroup and group:
        meter = equipment.find_one(
            {
                "PV": PV,
                "lgroup": lgroup,
                "group": group,
                "type": "sun",
                "main_sun": 1
            }
        )
    if meter:
        ID = str(meter["_id"])
    else:
        ID = ""
    return ID

#PR跟meterID
def PR_and_meter_ID(pv, PV):
    equipment = pv["equipment"]
    meter = equipment.find_one(
        {
            "PV": PV,
            "type": "pv_group"
        }
    )
    ID = str(meter["_id"])
    return ID

#日照計的data list
def irrh_cal(pv, ID, times, time_interval, round_number=1):
    irrh_data = []
    irrh_cal = pv["irrh_cal"]
    for time in times:
        if len(ID) > 0:
            irrh = irrh_cal.find_one(
                {
                    "ID": ID,
                    "time": time, 
                    "time_interval": time_interval
                }
            )
            if irrh:
                if irrh["irrh"]:
                    if type(round_number) == int:
                        irrh_data.append(round(irrh["irrh"], round_number))
                    else:
                        round_number = round(round_number)
                        irrh_data.append(round(irrh["irrh"], round_number))
                else:
                    irrh_data.append("---")
            else:
                irrh_data.append("---")
        else:
            irrh_data.append("---")
    return irrh_data

#meter的data list
def meter_cal(pv, ID, times, time_interval, round_number=1):
    meter_data = []
    meter_cal = pv["meter_cal"]
    for time in times:
        meter = meter_cal.find_one(
            {
                "ID": ID,
                "time": time, 
                "time_interval": time_interval
            }
        )
        if meter:
            if meter["kwh"]:
                if type(round_number) == int:
                    meter_data.append(round(meter["kwh"], round_number))
                else:
                    round_number = round(round_number)
                    meter_data.append(round(meter["kwh"], round_number))
            else:
                meter_data.append("---")
        else:
            meter_data.append("---")
    return meter_data

#pr的data list
def pr_cal(pv, ID, times, time_interval, round_number=1):
    pr_data = []
    pr_cal = pv["pr_cal"]
    for time in times:
        pr = pr_cal.find_one(
            {
                "ID": ID,
                "time": time, 
                "time_interval": time_interval
            }
        )
        if pr:
            if pr["pr"]:
                if type(round_number) == int:
                    pr_data.append(round(pr["pr"], round_number))
                else:
                    round_number = round(round_number)
                    pr_data.append(round(pr["pr"], round_number))
            else:
                pr_data.append("---")
        else:
            pr_data.append("---")
    return pr_data

#判斷是否為閏年並返回天數
def leap_year(year):
    if (year % 4) == 0:
        if (year % 100) == 0:
            if (year % 400) == 0:
                month_day = 29
            else:
                month_day = 28
        else:
            month_day = 29
    else:
        month_day = 28
    return month_day

#所有需要的時間
def set_time_interval(start_time, end_time, time_interval):
    during = []
    during2 = []
    if time_interval == "year":
        allyear_len = 0
        if start_time <= end_time:
            end_time += relativedelta(years=1)
            allyear_len = end_time.year - start_time.year

        allyear_len = int(allyear_len)
        deltatime = relativedelta(years=+1)

        start = date(start_time.year, 1, 1)
        nexttime = datetime.combine(start, time.min)

        for _ in range(allyear_len):
            during.append(nexttime)
            # during2.append(str(nexttime))
            nexttime += deltatime
    elif time_interval == "month":
        year_len = 0
        if start_time <= end_time:
            end_time += relativedelta(months=1)
            year_len = (end_time.year - start_time.year)*12 + end_time.month - start_time.month
        
        year_len = int(year_len)
        deltatime = relativedelta(months=+1)

        start = date(start_time.year, start_time.month, 1)
        nexttime = datetime.combine(start, time.min)

        for _ in range(year_len):
            during.append(nexttime)
            # during2.append(str(nexttime))
            nexttime += deltatime
    elif time_interval == "day":
        month_len = 0
        if start_time <= end_time:
            end_time += relativedelta(days=1)
            month_len = (end_time - start_time).total_seconds() / 86400
        
        month_len = int(month_len)
        deltatime = relativedelta(days=+1)

        start = date(start_time.year, start_time.month, start_time.day)
        nexttime = datetime.combine(start, time.min)

        for _ in range(month_len):
            during.append(nexttime)
            # during2.append(str(nexttime))
            nexttime += deltatime
    elif time_interval == "hour":
        day_len = 0
        if start_time <= end_time:
            start_time += relativedelta(hour=0, minute=0, second=0)
            end_time += relativedelta(days=+1, hour=0, minute=0, second=0)
            day_len = (end_time - start_time).total_seconds() / (86400/24)
        
        day_len = int(day_len)
        deltatime = relativedelta(hours=+1)

        nexttime = datetime.combine(start_time, time.min)

        for _ in range(day_len):
            during.append(nexttime)
            # during2.append(str(nexttime))
            nexttime += deltatime
    elif time_interval == "15min":
        day_len = 0
        if start_time <= end_time:
            start_time += relativedelta(hour=0, minute=0, second=0)
            end_time += relativedelta(days=+1, hour=0, minute=0, second=0)
            day_len = (end_time - start_time).total_seconds() / (60*15)
        
        day_len = int(day_len)
        deltatime = relativedelta(minutes=+15)

        nexttime = datetime.combine(start_time, time.min)

        for _ in range(day_len):
            during.append(nexttime)
            # during2.append(str(nexttime))
            nexttime += deltatime
    # return during, during2
    return during
#頁首資訊
def company_imformation(pv, object_id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(object_id)
        }
    )
    company_name = plant.get('client_info', {}).get('unit', '')
    tel = plant.get('client_info', {}).get('TEL', '')
    # company_name = plant.get("name", "")
    imformation = []
    imformation.append(company_name)
    imformation.append(tel)
    return imformation

#頁尾資訊
def company_imformation_footer(pv, object_id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(object_id)
        }
    )
    company_name = plant.get('paets_info', {}).get('unit', '')
    return company_name

#專案名稱
def project_name(pv, id, start_time, end_time, time_interval):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    equipment = pv["equipment"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    if plant and not equipment:
        name = plant["name"]
    elif not plant and equipment:
        PV_str = equipment["PV"]
        if equipment["type"] == "pv_lgroup":
            lgroup = equipment["name"]
            group = None
        elif equipment["type"] == "pv_group":
            lgroup = equipment["lgroup"]
            group = equipment["name"]

        if lgroup and group:
            name = PV_str + "-" + lgroup + "-" + group
        elif lgroup and not group:
            name = PV_str + "-" + lgroup
        else:
            name = PV_str
    return name+"_"+str(start_time)+"_"+str(end_time)+"_"+time_interval

def imformation_data(project_name, date, position, capacity, max_val, avg_val="0"):
    if max_val == "0" and avg_val == "0":
        imformation = {
            "專案名稱": project_name,
            "日期": date,
            "案場位置": position,
            "設置容量": capacity + "kW",
            "最大功率": "---",
            # "平均功率": "---"
        }
    else:
        imformation = {
            "專案名稱": project_name,
            "日期": date,
            "案場位置": position,
            "設置容量": capacity + "kW",
            "最大功率": max_val + "kW",
            # "平均功率": avg_val + "kW"
        }
    return imformation

def avg_max_value(pv, ID, collection, start_time, end_time, round_number=1, avg_status=False):
    max = 0
    if len(collection) > 0:
        data_collection = pv[collection]
        end_time = str(end_time).split(" ")[0] + " 23:59:59"
        end_time = datetime.strptime(end_time, '%Y-%m-%d %H:%M:%S')
        max_value = data_collection.find(
            {
                "ID": ID,
                "time":{"$gte": start_time, "$lt": end_time}
            }
        ).sort("p",-1).limit(1)
        if avg_status:
            avg_data = data_collection.find(
                {
                    "ID": ID,
                    "time": {"$gte": start_time, "$lt": end_time},
                    "p": {"$nin": [None, 0]}
                }
            )
            avg_list = []
            max = 0
            if avg_data:
                for average in avg_data:
                    avg_list.append(average["p"])
            if len(avg_list) > 0:
                avg = round(sum(avg_list)/len(avg_list), round_number)
            else:
                avg = 0
        else:
            avg = 0

        if max_value:
            for maximum in max_value:
                if maximum.get("p", 0) != None:
                    max = round(maximum.get("p", 0), round_number)
                else:
                    max = 0
    else:
        max = 0
        avg = 0
    if avg_status:
        return max, avg
    else:
        return max

#---------------------------報表製作-----------------------------------------------------

#初始設定
def document_initial(doc, margin=None):
    #設定邊界
    section = doc.sections[0]
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    if margin != None:
        section.left_margin = Cm(margin["left"])
        section.right_margin = Cm(margin["right"])
        section.top_margin = Cm(margin["top"])
        section.bottom_margin = Cm(margin["bottom"])
    #設定字型
    doc.styles['Normal'].font.size = Pt(12)
    #設定字體(英文)
    doc.styles['Normal'].font.name = 'Times New Roman'
    #中文
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

#資料表格
def add_table(doc, head_datas, datas, width, style=None):
    table = doc.add_table(rows=1, cols=len(head_datas), style=("Table Grid" if style is None else style))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_cells = table.rows[0].cells
    #表格標頭文字
    for index, head_item in enumerate(head_datas):
        head_cells[index].text = head_item
        head_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        head_cells[index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #表格資料
    for data in datas:
        row_cells = table.add_row().cells
        for index, cell in enumerate(row_cells):
            cell.text = str(data[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # #表格內寬度
    # for index, column in enumerate(table.columns):
    #     for cell in column.cells:
    #         cell.width = Cm(width[index])

    #表格內高度
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(0.6)
        else:
            row.height = Cm(0.64)

#刪除段落、圖片等
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    #p._p = p._element = None
    paragraph._p = paragraph._element = None

#資訊欄
def imformation_table(doc, imformation_data, time_interval):
    width = [3.9, 15.1]
    rows = 0
    index = 0
    keys = []
    row_index = 0
    key_index = 0

    #設定標題
    heading = doc.add_paragraph("")
    heading.style = doc.styles['Heading 1']
    #段前間距
    heading.paragraph_format.space_before = Pt(0)
    #段後間距
    heading.paragraph_format.space_after = Pt(15)

    title = ""
    if time_interval == "15min":
        title = "太陽光電系統維運每15分鐘報表"
    elif time_interval == "hour":
        title = "太陽光電系統維運每小時報表"
    elif time_interval == "day":
        title = "太陽光電系統維運每日報表"
    elif time_interval == "month":
        title = "太陽光電系統維運每月報表"
    elif time_interval == "year":
        title = "太陽光電系統維運每年報表"
    run = heading.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    
    # run.space_before = Pt(9)
    # run.space_after = Pt(9)
    # run.line_spacing = 1.5

    for value in imformation_data.values():
        if type(value) == list:
            rows += len(value)
        elif type(value) == str:
            rows += 1
    
    table = doc.add_table(rows=rows, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #table.cell(0, 0).text = "表角"

    for key in imformation_data.keys():
        keys.append(key)

    while row_index < rows:
        if type(imformation_data[keys[key_index]]) == list:
            for content in imformation_data[keys[key_index]]:
                row = table.rows[row_index]
                row.cells[0].text = keys[key_index]
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row.cells[1].text = content
                row_index += 1
        else:
            row = table.rows[row_index]
            row.cells[0].text = keys[key_index]
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row.cells[1].text = imformation_data[keys[key_index]]
            row_index += 1
        key_index += 1

    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])


    #合併重複名稱
    col = table.columns[0]
    row_index = 0
    while row_index < rows-1:
        if col.cells[row_index].text == col.cells[row_index+1].text:
            content = col.cells[row_index].text
            col.cells[row_index].merge(col.cells[row_index+1])
            col.cells[row_index].text = content
            col.cells[row_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            col.cells[row_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_index += 1

#設定頁首頁尾
def header(doc, imformations, logo, footer_info):
    section = doc.sections[0]
    header = section.header
    # paragraph = header.paragraphs[0]
    paragraph = header.add_paragraph()
    #段落間距
    paragraph.paragraph_format.space_after = Pt(0)
    #單行間距(長度為絕對距離，浮點數為間距)
    paragraph.paragraph_format.line_spacing = 1.0
    #刪除首段
    delete_paragraph(header.paragraphs[0])
    tab_stops = paragraph.paragraph_format.tab_stops
    logo_run = paragraph.add_run()


    # try:   # For logo not found
    #     logo_run.add_picture(logo, height=Inches(0.65))
    # except:
    #     pass
    text_run = paragraph.add_run('\t')
    text_run = paragraph.add_run(imformations[0])
    #設置右邊文字位置tab
    margin_end = Inches(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
    tab_stop = tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
    #字體大小、斜體
    text_run.font.size = Pt(14)
    text_run.font.italic = True

    for index, imformation in enumerate(imformations):
        if index == 0:
            #paragraph = header.paragraphs[0]
            continue
        else:
            paragraph = header.add_paragraph()
            #pass
        run = paragraph.add_run(""+imformation)
        if index == 0:
            #run.font.size = Pt(14)
            continue
        else:
            run.font.size = Pt(10)
        run.font.italic = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    #頁尾
    footer = section.footer
    # paragraph = header.paragraphs[0]
    paragraph = footer.add_paragraph()
    #段落間距
    paragraph.paragraph_format.space_after = Pt(0)
    #單行間距(長度為絕對距離，浮點數為間距)
    paragraph.paragraph_format.line_spacing = 1.0
    #刪除首段
    delete_paragraph(footer.paragraphs[0])
    tab_stops = paragraph.paragraph_format.tab_stops

    text_run = paragraph.add_run('\t')
    text_run = paragraph.add_run(footer_info)
    #設置右邊文字位置tab
    margin_end = Inches(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
    tab_stop = tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
    #字體大小、斜體
    text_run.font.size = Pt(14)
    text_run.font.italic = True
    
def table_data(dates, irrh_data, meter_data, pr_data):
    data = []
    for date, irrh, meter, pr in zip(dates, irrh_data, meter_data, pr_data):
        data.append([str(date).replace("-", "/"), irrh, meter, pr])
    return data

def report(name, header_imformation, imformation, table_head_datas, table_datas, static_path, time_interval, footer_info):
    logo = static_path + "/images/logo.jpg"

    #總長約19cm
    width = [4.71, 5.49, 4.4, 4.4]

    doc = Document()
    #初始設定
    document_initial(doc)
    #設定頁首
    header(doc, header_imformation, logo, footer_info)
    #資訊欄
    imformation_table(doc, imformation, time_interval)
    p = doc.add_paragraph()
    #資料表格
    add_table(doc, table_head_datas, table_datas, width=width)
    doc_name = name + ".docx"
    #doc.save(doc_name)
    doc_name =  quote(doc_name)

    # Create in-memory buffer
    file_stream = BytesIO()
    # Save the .docx to the buffer
    doc.save(file_stream)
    # Reset the buffer's file-pointer to the beginning of the file
    file_stream.seek(0)
    rv = send_file(
        file_stream,
        as_attachment=True,
        attachment_filename=doc_name,
        download_name=doc_name
    )
    return rv

#---------------------------派工工單報表-----------------------------------------------------
def dispatch_information_table(doc, imformation_data, title, width_set=None):
    width = [3.9, 15.1]
    if width_set:
        width = width_set
    rows = 0
    index = 0
    keys = []
    row_index = 0
    key_index = 0

    if title != "":
        #設定標題
        heading = doc.add_paragraph("")
        heading.style = doc.styles['Heading 1']
        #段前間距
        heading.paragraph_format.space_before = Pt(0)
        #段後間距
        heading.paragraph_format.space_after = Pt(15)
        
        run = heading.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    
    # run.space_before = Pt(9)
    # run.space_after = Pt(9)
    # run.line_spacing = 1.5

    for key in imformation_data.keys():
        keys.append(key)
        if type(imformation_data[key]) == list and key != "相片":
            rows += len(imformation_data[key])
        elif type(imformation_data[key]) == str:
            rows += 1
        elif key == "相片":
            rows += 1

    # for value in imformation_data.values():
    #     if type(value) == list:
    #         rows += len(value)
    #     elif type(value) == str:
    #         rows += 1
    
    table = doc.add_table(rows=rows, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #table.cell(0, 0).text = "表角"

    while row_index < rows:
        if type(imformation_data[keys[key_index]]) == list and keys[key_index] != "相片":
            for content in imformation_data[keys[key_index]]:
                row = table.rows[row_index]
                row.cells[0].text = keys[key_index]
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row.cells[1].text = content
                row_index += 1
        elif keys[key_index] == "相片":
            row = table.rows[row_index]
            row.cells[0].text = keys[key_index]
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = row.cells[1].paragraphs[0]
            run = paragraph.add_run()
            for picture in imformation_data[keys[key_index]]:
                try:
                    run.add_picture(picture, height=Inches(0.75))
                    run = paragraph.add_run("  ")
                    run = paragraph.add_run()
                except Exception as e:
                    print(exception_detail(e))
                    continue
            row_index += 1
        else:
            if  keys[key_index] not in ["原因分析", "項目名稱", "警報名稱"]:
                row = table.rows[row_index]
                row.cells[0].text = keys[key_index]
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row.cells[1].text = imformation_data[keys[key_index]]
            else:
                row = table.rows[row_index]
                paragraph = row.cells[0].paragraphs[0]
                run = paragraph.add_run(keys[key_index])
                run.bold = True
                paragraph = row.cells[1].paragraphs[0]
                run = paragraph.add_run(imformation_data[keys[key_index]])
                run.bold = True
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            row_index += 1

        key_index += 1

    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])


    #合併重複名稱
    col = table.columns[0]
    row_index = 0
    while row_index < rows-1:
        if col.cells[row_index].text == col.cells[row_index+1].text:
            content = col.cells[row_index].text
            col.cells[row_index].merge(col.cells[row_index+1])
            col.cells[row_index].text = content
            col.cells[row_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            col.cells[row_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_index += 1

def get_alarm_cause(db, dispatch_ID, alarm_list):
    return_list = []
    for alarm_dict in alarm_list:
        try:
            for alarm in db.alarm.find({'_id': ObjectId(alarm_dict['ID'])}):
                alarm['dispatch_ID'] = dispatch_ID
                alarm['_id'] = str(alarm['_id'])
                #get equip data
                equip_type = ''
                equip_name = ''
                system_translate = {"DG": "DG","PV": "地面型","BESS": "屋頂型","WT": "水面型"}
                type_translate = {"inv": "變流器", "string": "串電流錶", "io": "開關", "sun": "日照計", 
                "temp": "溫度計", "wind": "風速計", "meter": "智慧電錶", "pv_meter": "智慧電錶"}
                _data = db.plant.find_one({'_id': ObjectId(alarm['ID'])})
                if _data != None:
                    equip_type = '案場'
                    equip_name = _data.get('name', '')
                else:
                    _data = db.iot.find_one({'_id': ObjectId(alarm['ID'])})
                    if _data != None:
                        equip_type = '資料收集器'
                        equip_name = _data.get('name',  '資料收集器')
                    else:
                        _data = db.equipment.find_one({'_id': ObjectId(alarm['ID'])})
                        equip_name = _data.get('name',  '')
                        if _data.get('type') == 'pv_lgroup':
                            equip_type = '分區' if _data.get('Devide_type', '') == '' else type_translate.get(_data.get('Device_type', ''), _data.get('Device_type', ''))
                        elif _data.get('type') == 'pv_group':
                            equip_type = '分組' if _data.get('Devide_type', '') == '' else type_translate.get(_data.get('Device_type', ''), _data.get('Device_type', ''))
                        else:
                            equip_type = type_translate.get(_data.get('type', ''), _data.get('type', ''))
                
                alarm['equip_name'] = equip_name
                alarm['equip_type'] = equip_type
                alarm['cause_data'] = {
                    'ai_analysis': [],
                    'user_select': []
                }
                for cause_ID in alarm_dict.get('cause', {}):
                    #print(cause_ID)
                    # photo to pathname
                    alarm_dict['cause'][cause_ID]['photo_data'] = []
                    for filename in alarm_dict['cause'][cause_ID].get('photo', []):
                        alarm_dict['cause'][cause_ID]['photo_data'].append({
                            'filename': filename,
                            'filepath': "solar_static/dispatch/{}/{}/{}/{}".format(
                                dispatch_ID, alarm['_id'], cause_ID, filename
                            ) 
                        })
                    if cause_ID == "else":    #其他事項
                        alarm_dict['cause']['else']['group'] = '其他'
                        alarm_dict['cause']['else']['event'] = '其他'
                        alarm['cause_data']['user_select'].append(alarm_dict['cause']['else'])
                        continue   # Skip for searching db
                    for cause in db.alarm_cause.find({'_id': ObjectId(cause_ID)}):
                        #print(cause)
                        alarm_dict['cause'][cause_ID]['group'] = cause['alarm_group']
                        alarm_dict['cause'][cause_ID]['event'] = cause['event']
                        alarm['cause_data']['user_select'].append(alarm_dict['cause'][cause_ID])
                # sorting 
                # place else in last place
                new_user_select = []
                for cause in alarm['cause_data']['user_select']:
                    if cause['group'] == '其他':
                        new_user_select.append(cause)
                    else:
                        _new_user_select = [cause]
                        _new_user_select.extend(new_user_select)
                        new_user_select = _new_user_select
                alarm['cause_data']['user_select'] = new_user_select
                #print(alarm)
                #time fix
                try:
                    alarm['time'] = datetime.strftime(alarm['time'], '%Y-%m-%d %H:%M:%S')
                except:
                    pass
                try:
                    alarm['returntime'] = datetime.strftime(alarm['returntime'], '%Y-%m-%d %H:%M:%S')
                except:
                    alarm['returntime'] = '---'
                    pass
                return_list.append(alarm)
        except Exception as e:
            print(e)
            pass
    return return_list

def get_regular_wash_list(dispatch_ID, group_list):
    for group in group_list:
        try:
            group['dispatch_ID'] = dispatch_ID
            for child_name in group['child']:
                # photo to pathname
                group['child'][child_name]['photo_data'] = []
                for filename in group['child'][child_name].get('photo', []):
                    group['child'][child_name]['photo_data'].append({
                        'filename': filename,
                        'filepath': "solar_static/dispatch/{}/{}/{}/{}".format(
                            dispatch_ID, group['name'], child_name, filename
                        ) 
                    })
        except Exception as e:
            print(e)
    return group_list

def get_dispatch_data(db, request_dict):
    try:
        ID = request_dict['ID']
        ObjectId(ID)
    except Exception as e:
        print(e)

    #get dispatch data
    return_dict = {}
    for dispatch in db.dispatch.find({'_id': ObjectId(ID)}):
        dispatch['_id'] = str(dispatch['_id'])
        #get station data
        dispatch['station_data'] = {}
        for equip in db.equipment.find({'_id': ObjectId(dispatch['ID'])}):
            equip['_id'] = str(equip['_id'])
            equip['plant_data'] = {}
            for plant in db.plant.find({'name': equip.get('PV', '')}):
                plant['_id'] = str(plant['_id'])
                if plant.get('start_date', None) != None:
                    plant['start_date'] = datetime.strftime(plant['start_date'], '%Y-%m-%d')
                equip['plant_data'] = plant
            dispatch['station_data'] = equip
        #time strftime
        if dispatch.get('dispatch_time', None) != None:
            dispatch['dispatch_time'] = datetime.strftime(dispatch['dispatch_time'], '%Y-%m-%d')
        #maintainer data
        maintainer_ID = dispatch.get('maintainer_ID', None)
        dispatch['maintainer_data'] = {
            'name': '',
            'tel': ''
        }
        if maintainer_ID != None:
            maintainer_data = db.users.find_one({'_id': ObjectId(maintainer_ID)})
            if 'user_data' in maintainer_data:
                dispatch['maintainer_data'] = maintainer_data['user_data']
        #alarm_data
        if 'alarm' in dispatch['type']:
            dispatch['alarm_list'] = get_alarm_cause(db, dispatch['_id'], dispatch.get('alarm_list', []))
        #regular_data
        if 'regular' in  dispatch['type']:
            dispatch['regular_list'] = get_regular_wash_list(dispatch['_id'], dispatch.get('regular_list', []))
        #wash_data
        if 'wash' in dispatch['type']:
            dispatch['wash_list'] = get_regular_wash_list(dispatch['_id'], dispatch.get('wash_list', []))
        #print(dispatch)
        return_dict = dispatch
    return return_dict

def report_dispatch_regular(doc, dispatch_data):
    regular_list = dispatch_data.get("regular_list", [])
    if len(regular_list) > 0:
        p = doc.add_paragraph()

        #設定標題
        heading = doc.add_paragraph("")
        heading.style = doc.styles['Heading 1']
        #段前間距
        heading.paragraph_format.space_before = Pt(0)
        #段後間距
        heading.paragraph_format.space_after = Pt(15)
        
        run = heading.add_run("工單內容-定檢")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

        for regular_data in regular_list:
            p = doc.add_paragraph()
            #段前間距
            p.paragraph_format.space_before = Pt(0)
            #段後間距
            # p.paragraph_format.space_after = Pt(15)
            run = p.add_run("群組名稱: {}".format(regular_data.get("name", "---")))
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

            child_keys = regular_data.get("child", {}).keys()
            events = regular_data.get("child", {})

            result = "---"
            if len(child_keys) > 0:
                for child_key in child_keys:
                    event = events[child_key]
                    event_name = event.get("name", "---")
                    category_type = event.get("category", None)
                    if category_type == "choice":
                        category_type = "正常/異常"
                    elif category_type == "numeric":
                        category_type = "數值"
                    
                    table_infomation = {
                        "項目名稱": event_name,
                        "類別": category_type if category_type else "---"
                    }

                    if category_type != None:
                        if category_type == "正常/異常":
                            result = event.get("choice", "---")
                            table_infomation["檢測結果"] = result
                        elif category_type == "數值":
                            result = event.get("value", "---")
                            suggest_value = event.get("suggest_value", "---")
                            table_infomation["實際值/要求值"] = "{}/{}".format(result, suggest_value)
                    table_infomation["說明"] = event.get("info", "")

                    if len(event.get("photo_data", [])) > 0:
                        pictures = []
                        for photo_data in event["photo_data"]:
                            pictures.append(current_app.config['UPLOAD_FOLDER']+"/"+photo_data["filepath"].split("solar_static/")[-1])
                            # pictures.append("ntust.png")
                        table_infomation["相片"] = pictures
                    # print(table_infomation.get("相片", ""))
                    # else:
                    #     p = doc.add_paragraph()
                    #     add_table(doc, ["無相片"], [], 0)

                    dispatch_information_table(doc, table_infomation, "")
            p = doc.add_paragraph()
    elif len(regular_list) == 0:
        # p = doc.add_paragraph()
        # #設定標題
        # heading = doc.add_paragraph("")
        # heading.style = doc.styles['Heading 1']
        # #段前間距
        # heading.paragraph_format.space_before = Pt(0)
        # #段後間距
        # heading.paragraph_format.space_after = Pt(15)
        
        # run = heading.add_run("工單內容-定檢")
        # run.bold = True
        # run.font.color.rgb = RGBColor(0, 0, 0)
        # run.font.name = 'Times New Roman'
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        # add_table(doc, ["無定檢"], [], 0)
        return False
    else:
        return False

def report_dispatch_wash(doc, dispatch_data):
    wash_list = dispatch_data.get("wash_list", [])
    if len(wash_list) > 0:
        p = doc.add_paragraph()

        #設定標題
        heading = doc.add_paragraph("")
        heading.style = doc.styles['Heading 1']
        #段前間距
        heading.paragraph_format.space_before = Pt(0)
        #段後間距
        heading.paragraph_format.space_after = Pt(15)
        
        run = heading.add_run("工單內容-清洗")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

        for wash_data in wash_list:
            p = doc.add_paragraph()
            #段前間距
            p.paragraph_format.space_before = Pt(0)
            #段後間距
            # p.paragraph_format.space_after = Pt(15)
            run = p.add_run("群組名稱: {}".format(wash_data.get("name", "---")))
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

            child_keys = wash_data.get("child", {}).keys()
            events = wash_data.get("child", {})

            result = "---"
            if len(child_keys) > 0:
                for child_key in child_keys:
                    event = events[child_key]
                    event_name = event.get("name", "---")
                    category_type = event.get("category", None)
                    if category_type == "choice":
                        category_type = "正常/異常"
                    elif category_type == "numeric":
                        category_type = "數值"
                    
                    table_infomation = {
                        "項目名稱": event_name,
                        "類別": category_type if category_type else "---"
                    }

                    if category_type != None:
                        if category_type == "正常/異常":
                            result = event.get("choice", "---")
                            table_infomation["檢測結果"] = result
                        elif category_type == "數值":
                            result = event.get("value", "---")
                            suggest_value = event.get("suggest_value", "---")
                            table_infomation["實際值/要求值"] = "{}/{}".format(result, suggest_value)
                    table_infomation["說明"] = event.get("info", "")

                    if len(event.get("photo_data", [])) > 0:
                        pictures = []
                        for photo_data in event["photo_data"]:
                            # pictures.append("ntust.png")
                            pictures.append(current_app.config['UPLOAD_FOLDER']+"/"+photo_data["filepath"].split("solar_static/")[-1])
                            
                        table_infomation["相片"] = pictures
                    # print(table_infomation.get("相片", ""))
                    # else:
                    #     p = doc.add_paragraph()
                    #     add_table(doc, ["無相片"], [], 0)

                    dispatch_information_table(doc, table_infomation, "")
            p = doc.add_paragraph()
    elif len(wash_list) == 0:
        # p = doc.add_paragraph()
        # #設定標題
        # heading = doc.add_paragraph("")
        # heading.style = doc.styles['Heading 1']
        # #段前間距
        # heading.paragraph_format.space_before = Pt(0)
        # #段後間距
        # heading.paragraph_format.space_after = Pt(15)
        
        # run = heading.add_run("工單內容-清洗")
        # run.bold = True
        # run.font.color.rgb = RGBColor(0, 0, 0)
        # run.font.name = 'Times New Roman'
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        # add_table(doc, ["無清洗"], [], 0)
        return False
    else:
        return False

def report_dispatch_alarm(doc, dispatch_data):
    alarm_list = dispatch_data.get("alarm_list", [])
    if len(alarm_list) > 0:
        p = doc.add_paragraph()

        #設定標題
        heading = doc.add_paragraph("")
        heading.style = doc.styles['Heading 1']
        #段前間距
        heading.paragraph_format.space_before = Pt(0)
        #段後間距
        heading.paragraph_format.space_after = Pt(15)
        
        run = heading.add_run("工單內容-告警檢修")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

        for alarm_data in alarm_list:
            event = alarm_data.get("event", "---")
            station_data = dispatch_data.get("station_data", {})
            station = "{}/{}".format(station_data.get("name", ""), station_data.get("PV", ""))
            equip_name = alarm_data.get("equip_name", "---")
            time = alarm_data.get("time", "---")
            return_time = alarm_data.get("returntime", "---")
            table_infomation = {
                "警報名稱": event,
                "廠區": station,
                "設備名稱": equip_name,
                "警報發生": time,
                "警報修復": return_time
            }
            dispatch_information_table(doc, table_infomation, "")
            user_selects = alarm_data.get("cause_data", {}).get("user_select", [])
            if len(user_selects)>0:
                for user_select in user_selects:
                    cause_name = "{}/{}".format(user_select.get("group", ""), user_select.get("event", ""))
                    fix = user_select.get("fix", False)
                    info = user_select.get("info", "")
                    photos = []
                    for photo in user_select.get("photo_data", []):
                        # photos.append("ntust.png")
                        photos.append(current_app.config['UPLOAD_FOLDER']+"/"+photo["filepath"].split("solar_static/")[-1])
                    cause_dict = {
                        "原因分析": cause_name,
                        "修復/未修復": "修復" if fix == True else "未修復",
                        "說明": info
                    }
                    if len(photos)>0:
                        cause_dict["相片"] = photos
                    dispatch_information_table(doc, cause_dict, "")

            p = doc.add_paragraph()
    elif len(alarm_list) == 0:
        # p = doc.add_paragraph()
        # #設定標題
        # heading = doc.add_paragraph("")
        # heading.style = doc.styles['Heading 1']
        # #段前間距
        # heading.paragraph_format.space_before = Pt(0)
        # #段後間距
        # heading.paragraph_format.space_after = Pt(15)
        
        # run = heading.add_run("工單內容-告警檢修")
        # run.bold = True
        # run.font.color.rgb = RGBColor(0, 0, 0)
        # run.font.name = 'Times New Roman'
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        # add_table(doc, ["無告警"], [], 0)
        return False
    else:
        return False

def report_dispatch_work(doc, dispatch_data):
    working_data = dispatch_data.get("working_data", [])
    working_table_dict = {}
    if len(working_data)>0:
        working_data = working_data[-1]
        working_table_dict["工作時數"] = "{} 小時".format(working_data.get("working_hour", "---"))
    working_table_dict["本次維運總成本 (元)"] = "{}".format(dispatch_data.get("auto_review_cost", "---"))
    detail = dispatch_data.get("auto_review_cost_detail", {})
    predict_dispatch_cost = detail.get("predict_dispatch_cost", None)
    transit_fee = detail.get("transit_fee", None)
    worker_fee =  detail.get("worker_fee", None)
    detail_name = ""
    if predict_dispatch_cost != None:
        detail_name += "硬體成本 {}".format(predict_dispatch_cost)
    if transit_fee != None:
        if len(detail_name)>0:
            detail_name += " + 交通成本 {}".format(transit_fee)
        else:
            detail_name += "交通成本 {}".format(transit_fee)
    if worker_fee != None:
        if len(detail_name)>0:
            detail_name += " + 人事成本 {}".format(worker_fee)
        else:
            detail_name += "人事成本 {}".format(worker_fee)
    working_table_dict["細節"] = detail_name

    p = doc.add_paragraph()
    dispatch_information_table(doc, working_table_dict, "成本", width_set=[4.2, 14.8])

def report_dispatch(db, request_dict):
    try:
        dispatch_data = get_dispatch_data(db, request_dict)
        print(dispatch_data)

        stage = {
            "wait_for_priority": "等待優先度排序",
            "wait_for_take": "待接單",
            "merged": "已合併",
            "took_wait_date_enter": "等待輸入派工日期",
            "wait_for_dispatch": "等待派工",
            "wait_admin_confirm_date": "等待管理人員確認派工日期",
            "dispatched_wait_for_review": "等待AI驗收",
            "auto_reviewed_wait_for_manual": "等待管理人員協助驗收",
            "review_failed": "驗收失敗",
            "dispatch_finish": "已完成工單"
        }

        dispatch_infomation = {
            "單號": dispatch_data.get("name", ""),
            "派工日期": dispatch_data.get("dispatch_time", ""),
            "狀態": stage[dispatch_data.get("stage", "")],
            "預估維運成本": str(dispatch_data.get("predict_dispatch_cost", "")),
            "維運人員": [dispatch_data.get("maintainer_data", {}).get("name", ""), dispatch_data.get("maintainer_data", {}).get("tel", "")]
        }
        # print(dispatch_infomation)

        station_infomation = {
            "掛錶日期": dispatch_data.get("station_data", {}).get("plant_data", {}).get("start_date", ""),
            "裝置容量": str(dispatch_data.get("station_data", {}).get("capacity", "---"))+" kWp",
            "案場位置": dispatch_data.get("station_data", {}).get("plant_data", {}).get("plant_address", ""),
            "經緯度": dispatch_data.get("station_data", {}).get("plant_data", {}).get("coordinates", ""),
            "模組型號": dispatch_data.get("station_data", {}).get("module_model", ""),
            "備註欄": ["{}. {}".format(index+1, event) for index, event in enumerate(dispatch_data.get("station_data", {}).get("plant_data", {}).get("event", []))],
            "受託者": [dispatch_data.get("station_data", {}).get("plant_data", {}).get("client_info", {}).get("unit", ""), dispatch_data.get("station_data", {}).get("plant_data", {}).get("client_info", {}).get("TEL", "")],
        }
        # print(station_infomation)

        start_address = dispatch_data.get("working_data", [])[len(dispatch_data.get("working_data", []))-1].get("transit", {}).get("start_address", "")

        work_infomation = {
            "出發地址": start_address if start_address else "---",
            "預估交通成本":  "NTD$ "+str(dispatch_data.get("working_data", [])[len(dispatch_data.get("working_data", []))-1].get("transit", {}).get("fee", ""))
        }
        # print(work_infomation)

        doc = Document()
        #初始設定
        document_initial(doc)

        dispatch_information_table(doc, dispatch_infomation, title="工單資訊")

        p = doc.add_paragraph()
        dispatch_information_table(doc, station_infomation, title="案場資訊")

        p = doc.add_paragraph()
        dispatch_information_table(doc, work_infomation, title="工作路徑")

        report_dispatch_regular(doc, dispatch_data)

        report_dispatch_wash(doc, dispatch_data)

        report_dispatch_alarm(doc, dispatch_data)

        report_dispatch_work(doc, dispatch_data)

        filename = "派工報表-{}.docx".format(dispatch_data.get("name", ""))
        print(filename)
        doc.save("報表.docx")
        # doc.save(filename)
        doc_name = quote(filename)

        # Create in-memory buffer
        file_stream = BytesIO()
        # Save the .docx to the buffer
        doc.save(file_stream)
        # Reset the buffer's file-pointer to the beginning of the file
        file_stream.seek(0)
        rv = send_file(
            file_stream,
            as_attachment=True,
            attachment_filename=doc_name,
            download_name=doc_name
        )
        return rv
        # doc.save(filename)
    except Exception as e:
        print("docx_produce:")
        print(exception_detail(e))

#-------------------------------------------------------------------------------------------
# 茂鴻版本

def table_light_to_grid(table):
    for column in table.columns:
        for cell in column.cells:

            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.bold = False

            #GET CELLS XML ELEMENT
            cell_xml_element = cell._tc
            #RETRIEVE THE TABLE CELL PROPERTIES
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            #CREATE SHADING OBJECT
            shade_obj = OxmlElement('w:shd')
            #SET THE SHADING OBJECT
            shade_obj.set(qn('w:fill'), "fff")
            #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
            table_cell_properties.append(shade_obj)

def set_cell_border(table):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    kwargs = {
        "top": {"sz": 0, "val": "single", "color": "#000000"},
        "bottom": {"sz": 1, "color": "#000000", "val": "single"},
        "start": {"sz": 6, "color": "#000000", "val": "single"},
        "end": {"sz": 6, "color": "#000000", "val": "single"},
    }
    for column in table.columns:
        for cell in column.cells:
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.0
            #段前間距
            paragraph.paragraph_format.space_before = Pt(3)
            #段後間距
            paragraph.paragraph_format.space_after = Pt(3)

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # check for tag existnace, if none found, then create one
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)

            # list over all available tags
            for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
                edge_data = kwargs.get(edge)
                if edge_data:
                    tag = 'w:{}'.format(edge)

                    # check for tag existnace, if none found, then create one
                    element = tcBorders.find(qn(tag))
                    if element is None:
                        element = OxmlElement(tag)
                        tcBorders.append(element)

                    # looks like order of attributes is important
                    for key in ["sz", "val", "color", "space", "shadow"]:
                        if key in edge_data:
                            element.set(qn('w:{}'.format(key)), str(edge_data[key]))

#設定頁首
def dispatch_maohong_header(doc, logo):
    doc.sections[0].header_distance = Cm(0.0) 
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    header = section.header
    paragraph = header.paragraphs[0]
    # paragraph = header.add_paragraph()
    #段落間距
    # paragraph.paragraph_format.space_after = Pt(0)
    #單行間距(長度為絕對距離，浮點數為間距)
    # paragraph.paragraph_format.line_spacing = 1.0
    #刪除首段
    # delete_paragraph(header.paragraphs[0])
    logo_run = paragraph.add_run()

    try:   # For logo not found
        logo_run.add_picture(logo, height=Inches(0.5))
        # logo_run.add_picture(logo)
    except:
        print("no logo")

#進場紀錄表
def dispatch_maohong_record_information(doc, information_data, maintain_info):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 0.0

    width = [8.3, 10.6]
    rows = 2
    table = doc.add_table(rows=rows, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    data_list = [
        [
            "進場紀錄表", 
            "報告單號: {}".format(information_data["dispatch_number"])
        ],
        [
            "專案代號/案場名稱：\n{}".format(information_data["plant_name"]), 
            "進場日期：{}\n進場人員(廠商)：{}".format(information_data["entry_time"], information_data["worker"])
        ]
    ]

    for index, column in enumerate(table.columns):
        for y_index, cell in enumerate(column.cells):
            cell.width = Cm(width[index])

            paragraph = cell.paragraphs[0]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.add_run(data_list[y_index][index])
            run.font.size = Pt(16)
            if y_index == 0:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                if index == 0:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run.font.size = Pt(20)
            else:
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)


    width = [4.15, 4.15, 10.6]
    table = doc.add_table(rows=6, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])
    col = table.columns[0]
    for i in range(5):
        col.cells[0].merge(col.cells[i+1])
    paragraph = col.cells[0].paragraphs[0]
    run = paragraph.add_run("案場基本資料")
    run.font.size = Pt(16)
    col.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    col.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    col = table.columns[1]
    base_datas = ["裝置量(kWp)", "模組", "變流器", "型式/面向/角度", "掛錶日期", "上次巡檢日期"]
    for j, base_data in enumerate(base_datas):
        paragraph = col.cells[j].paragraphs[0]
        run = paragraph.add_run(base_data)
        run.font.size = Pt(16)
        col.cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        col.cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    col = table.columns[2]
    base_data_keys = ["capacity", "model", "inverter", "angle", "start_date", "last_date"]
    for j, base_data_key in enumerate(base_data_keys):
        paragraph = col.cells[j].paragraphs[0]
        run = paragraph.add_run(information_data[base_data_key])
        if j == 0:
            run.font.size = Pt(14)
        elif j ==1:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(16)
        col.cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        col.cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  
    width = [1.02, 17.88]
    table = doc.add_table(rows=2+len(maintain_info["alarm"]), cols=2, style="Normal Table")
    set_cell_border(table)
    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])
    col = table.columns[0]
    for i in range(1+len(maintain_info["alarm"])):
        col.cells[0].merge(col.cells[i+1])
    paragraph = col.cells[0].paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run("檢修保養內容")
    run.font.size = Pt(18)
    col.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    col.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    col = table.columns[1]
    cell = col.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(maintain_info["regular"])
    run.font.size = Pt(14)

    cell = col.cells[1]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(maintain_info["wash"])
    run.font.size = Pt(14)

    for k in range(len(maintain_info["alarm"])):
        cell = col.cells[k+2]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(maintain_info["alarm"][k])
        run.font.size = Pt(14)

    table = doc.add_table(rows=1, cols=1, style="Table Grid")
    for row in table.rows:
        row.height = Cm(7)
        for cell in row.cells:
            cell.width = Cm(18.9)
    paragraph = table.cell(0, 0).paragraphs[0]
    run = paragraph.add_run("本次進場的結果與後續處置：")
    run.font.size = Pt(18)

    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    table.cell(0, 0).height = Cm(1.98)
    table.cell(0, 0).width = Cm(9.45)
    paragraph = table.cell(0, 0).paragraphs[0]
    run = paragraph.add_run("主管簽名：")
    run.font.size = Pt(22)
    table.cell(0, 1).height = Cm(1.98)
    table.cell(0, 1).width = Cm(9.45)
    paragraph = table.cell(0, 1).paragraphs[0]
    run = paragraph.add_run("檢測人員簽名：")
    run.font.size = Pt(22)

def dispatch_last_dispatch_time(db, dispatch_data):
    ID = dispatch_data.get("ID", "")
    work_start_time = ""
    this_time = ""

    finish_time = dispatch_data["finish_time"]
    working_hour = dispatch_data.get("working_data", [])[-1]["working_hour"]
    this_time = finish_time - relativedelta(hours=int(working_hour), minute=0, second=0)
    this_time = datetime.strftime(this_time, "%Y-%m-%d")
    this_time = this_time.split("-")
    this_time = this_time[0]+"年"+this_time[1]+"月"+this_time[2]+"日"
    types = []
    for t in dispatch_data["type"]:
        types.append({"type": {"$all": [t]}})

    for data in db.dispatch.find({"ID": ID, "stage": "dispatch_finish", "$or": types}).sort("finish_time", -1).skip(1).limit(1):
        finish_time = data["finish_time"]
        working_hour = data.get("working_data", [])[-1]["working_hour"]
        work_start_time = finish_time - relativedelta(hours=int(working_hour), minute=0, second=0)
        work_start_time = datetime.strftime(work_start_time, "%Y-%m-%d")
        work_start_time = work_start_time.replace("-", "/")
    return work_start_time, this_time

def dispatch_find_module_inverter(db, dispatch_data):
    slices = dispatch_data.get("station_data", {}).get("plant_data", {}).get("module", "")

    PV = dispatch_data.get("station_data", {}).get("PV", "")
    lgroup = dispatch_data.get("station_data", {}).get("name", "")

    module_model = db.equipment.find_one({"PV": PV, "lgroup": lgroup, "type": "inv"}).get("module_model", "")
    Device_model = db.equipment.find_one({"PV": PV, "lgroup": lgroup, "type": "inv"}).get("Device_model", "")
    module_name = str(module_model)+" * "+str(slices)+" 片"
    inverter_number = db.equipment.count_documents({"PV": PV, "lgroup": lgroup, "type": "inv"})
    inverter_name = str(Device_model)+"*"+str(int(inverter_number))+" 台"
    return module_name, inverter_name

#檢修保養內容
def get_maintaince(db, dispatch_data):
    schedule_ID = dispatch_data.get("schedule_ID")
    types = dispatch_data.get("type", [])
    title_name = ""
    title_name_en = ""
    if type(schedule_ID) == str:
        schedule_datas = db.dispatch_schedule.find({"_id": ObjectId(schedule_ID)})
    elif type(schedule_ID) == list:
        filter = {"$or": []}
        for ID in schedule_ID:
            filter["$or"].append({"_id": ObjectId(ID)})
        schedule_datas = db.dispatch_schedule.find(filter)
    repeat_list = ["臨時", "不重複", "週檢", "月檢", "季檢", "半年檢", "年檢"]
    regular_name = ""
    if 'regular' in types:
        repeat = { 
            "none": "不重複",
            "week": "週檢",
            "month": "月檢",
            "quarter": "季檢",
            "year": "年檢",
            "temporary": "臨時",
            "half_year": "半年檢"
        }
        name_dict = {
            "不重複": "例行巡檢報表",
            "週檢": "例行巡檢周報表",
            "月檢": "例行巡檢月報表",
            "季檢": "例行巡檢季報表",
            "年檢": "例行巡檢年報表",
            "臨時": "例行巡檢報表",
            "半年檢": "例行巡檢半年報表"
        }
        name_dict_en = {
            "不重複": "ROUTINE INSPECTION REPORT",
            "週檢": "ROUTINE INSPECTION WEEK REPORT",
            "月檢": "ROUTINE INSPECTION MONTH REPORT",
            "季檢": "ROUTINE INSPECTION SEASON REPORT",
            "年檢": "ROUTINE INSPECTION YEAR REPORT",
            "臨時": "ROUTINE INSPECTION REPORT",
            "半年檢": "ROUTINE INSPECTION SEMI-ANNUAL REPORT"
        }
        if len(dispatch_data.get("repeat", [])) > 0:
            repeats = []
            for r in dispatch_data.get("repeat"):
                repeats.append(repeat.get(r))
        else:
            repeats = []
            for schedule_data in schedule_datas:
                repeats.append(repeat.get(schedule_data.get("repeat")))

        regular_name += "■定期檢查："
        for repeat_data in repeat_list:
            if repeat_data in repeats:
                regular_name += "■{} ".format(repeat_data)
                if title_name != "":
                    title_name += "、"
                    title_name_en += "\n"
                title_name += name_dict[repeat_data]
                title_name_en += name_dict_en[repeat_data]
            else:
                regular_name += "□{} ".format(repeat_data)
        regular_name += "□其他："
    else:
        regular_name = "□定期檢查：□臨時 □不重複 □週檢 □月檢 □季檢 □半年檢 □年檢 □其他："
    if 'wash' in types:
        wash_name = "■清潔保養：■除草 ■模組清潔 ■箱體保養 ■支架補漆"
        if len(title_name)>0:
            title_name += "、"
            title_name_en += "\n"
        title_name += "清潔報表"
        title_name_en += "WASH REPORT"
    else:
        wash_name = "□清潔保養：□除草 □模組清潔 □箱體保養 □支架補漆"

    alarm_causes = {}
    alarm_name = []
    _alarm_causes = {}
    for alarm_cause in db.alarm_cause.find({"show": 1}):
        if alarm_cause.get("alarm_group") not in alarm_causes:
            alarm_causes[alarm_cause.get("alarm_group")] = []
        alarm_causes[alarm_cause.get("alarm_group")].append(alarm_cause.get("event"))
    if 'alarm' in types:
        if len(title_name)>0:
            title_name += "、"
            title_name_en += "\n"
        title_name += "維修報表"
        title_name_en += "REPAIR REPORT"

        _alarm_name = "■異常排除："
        name = " ("
        for alarm in dispatch_data.get("alarm_list", []):
            for cause_data in alarm.get("cause_data", {}).get("user_select", []):
                if cause_data.get("group") not in _alarm_causes:
                    _alarm_causes[cause_data.get("group")] = []
                _alarm_causes[cause_data.get("group")].append(cause_data.get("event"))
        for key in alarm_causes.keys():
            for cause in alarm_causes[key]:
                if cause in _alarm_causes.get(key, []):
                    name += "■{} ".format(cause)
                else:
                    name += "□{} ".format(cause)
            if "■" in name:
                _alarm_name = _alarm_name+"■"+key+name+")"
            else:
                _alarm_name = _alarm_name+"□"+key+name+")"
            alarm_name.append(_alarm_name)
            _alarm_name = "\t\t  "
            name = "("

    else:
        _alarm_name = "□異常排除："
        for key in alarm_causes.keys():
            _alarm_name += "□{} (".format(key)
            for cause in alarm_causes[key]:
                _alarm_name += "□{} ".format(cause)
            _alarm_name += ")"
            alarm_name.append(_alarm_name)
            _alarm_name = "\t\t  "

    return regular_name, wash_name, alarm_name, title_name, title_name_en

#檢查項目
def list_table(doc, target_list, list_name):
    choice_list = ["y", "n", "NA"]
    normal_choice = {"y": "正常", "n": "異常", "NA": "N/A"}
    yn_choice = {"y": "是", "n": "否", "NA": "N/A"}
    left_list = []
    right_list = []     # 無子項目時長度為0
    if len(target_list) > 0 and list_name in ["wash_list", "regular_list"]:
        for target in target_list:
            left_list = []
            right_list = []
            title = target.get("number_name", "")   # ex: 1. 半年檢項目
            if title == "":
                title = target.get("name", "")   # ex: 1. 半年檢項目
            if title != "":
                left_list.append(title)
            if len(target.get("child", {}).keys())>0:
                right_list.append("")
                for key, child in target.get("child", {}).items():
                    right_name = ""
                    item_name = child.get("number_name", "")   # ex: 2-1 請填入數值
                    if item_name == "":
                        item_name = child.get("name", "")
                    if item_name != "":
                        left_list.append(item_name)
                        if child.get("category", "") == "numeric":
                            right_name += "數值：{} ".format(child.get("value"))
                            if child.get("suggest_value", None) != None:
                                right_name += " 建議數值：{} ".format(child.get("suggest_value"))
                        elif child.get("category", "") == "normal_choice_na":
                            choice = child.get("choice")
                            for choice_key in choice_list:
                                if choice == choice_key:
                                    right_name += "■{} ".format(normal_choice[choice])
                                else:
                                    right_name += "□{} ".format(normal_choice[choice_key])
                        elif child.get("category", "") == "yes_choice_na":
                            choice = child.get("choice")
                            for choice_key in choice_list:
                                if choice == choice_key:
                                    right_name += "■{} ".format(yn_choice[choice])
                                else:
                                    right_name += "□{} ".format(yn_choice[choice_key])
                        elif child.get("category", "") == "choice":
                            choice = child.get("choice")
                            if choice == "正常":
                                right_name += "■正常 □異常 "
                            else:
                                right_name += "□正常 ■異常 "
                        right_name += "說明：{}".format(child.get("info", ""))
                        right_list.append(right_name)

            if len(right_list) == 0:
                table = doc.add_table(rows=1, cols=1, style="Table Grid")
                paragraph = table.cell(0, 0).paragraphs[0]
                for index, left_data in enumerate(left_list):
                    run = paragraph.add_run(left_data)
                    run.font.size = Pt(12)
                    if index+1 < len(left_list):
                        run = paragraph.add_run("\n")
            else:
                table = doc.add_table(rows=1, cols=2, style="Table Grid")
                paragraph = table.cell(0, 0).paragraphs[0]
                for index, left_data in enumerate(left_list):
                    run = paragraph.add_run(left_data)
                    run.font.size = Pt(12)
                    if index+1 < len(left_list):
                        run = paragraph.add_run("\n")
                paragraph = table.cell(0, 1).paragraphs[0]
                for index, right_data in enumerate(right_list):
                    if "說明：" in right_data:
                        left = right_data.split("說明：")[0]
                        right = right_data.split("說明：")[1]
                        run = paragraph.add_run(left+"說明：")
                        run.font.size = Pt(12)
                        run = paragraph.add_run(right)
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        if index+1 < len(right_list):
                            run = paragraph.add_run("\n")
                    else:
                        run = paragraph.add_run(right_data)
                        if index+1 < len(right_list):
                            run = paragraph.add_run("\n")
                    run.font.size = Pt(12)
        doc.add_page_break()
    elif len(target_list) > 0 and list_name in ["alarm_list"]:
        for target in target_list:
            left_list = []
            right_list = []
            title = target.get("event", "")   # ex: 分組_資料收集器_通訊斷線
            if title != "":
                left_list.append(title)
            if len(target.get("cause_data", {}).get("user_select", []))>0:
                right_list.append("")
                for user_select in target.get("cause_data", {}).get("user_select", []):
                    left_name = ""
                    right_name = ""
                    left_name += user_select.get("group", "")+"/"+user_select.get("event", "")
                    if left_name != "/":
                        left_list.append(left_name)
                        if user_select.get("fix", False) == True:
                            right_name += "■修復 □未修復 "
                        else:
                            right_name += "□修復 ■未修復 "
                        right_name += "說明：{}".format(user_select.get("info", ""))
                        right_list.append(right_name)

            if len(right_list) == 0:
                table = doc.add_table(rows=1, cols=1, style="Table Grid")
                paragraph = table.cell(0, 0).paragraphs[0]
                for index, left_data in enumerate(left_list):
                    run = paragraph.add_run(left_data)
                    run.font.size = Pt(12)
                    if index+1 < len(left_list):
                        run = paragraph.add_run("\n")
            else:
                table = doc.add_table(rows=1, cols=2, style="Table Grid")
                paragraph = table.cell(0, 0).paragraphs[0]
                for index, left_data in enumerate(left_list):
                    run = paragraph.add_run(left_data)
                    run.font.size = Pt(12)
                    if index+1 < len(left_list):
                        run = paragraph.add_run("\n")
                paragraph = table.cell(0, 1).paragraphs[0]
                for index, right_data in enumerate(right_list):
                    if "說明：" in right_data:
                        left = right_data.split("說明：")[0]
                        right = right_data.split("說明：")[1]
                        run = paragraph.add_run(left+"說明：")
                        run.font.size = Pt(12)
                        run = paragraph.add_run(right)
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        if index+1 < len(right_list):
                            run = paragraph.add_run("\n")
                    else:
                        run = paragraph.add_run(right_data)
                        if index+1 < len(right_list):
                            run = paragraph.add_run("\n")
                    run.font.size = Pt(12)

        doc.add_page_break()

#相片
def list_picture(doc, target_list, list_name):
    title_list = []
    photo_dict_list = []
    photo_dict = {}
    if list_name in ["wash_list", "regular_list"]:
        for target in target_list:
            for child in target.get("child", {}).values():
                if len(child.get("photo_data", [])) > 0:
                    if target.get("name", "") not in title_list:
                        title_list.append(target.get("name", ""))
                        photo_dict_list.append(target)          # 留有相片的dict
                        break
        for index, data in enumerate(photo_dict_list):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(title_list[index].split(" ", 1)[-1])
            run.font.size = Pt(20)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            for child in data.get("child", {}).values():
                if len(child.get("photo_data", [])) > 0:
                    table = doc.add_table(rows=1, cols=1, style="Table Grid")
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    paragraph = table.cell(0, 0).paragraphs[0]
                    run = paragraph.add_run(child.get("name", "").split(" ", 1)[-1])
                    run.font.size = Pt(14)
                    table.cell(0, 0).width = Cm(18.9)
                    table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    table.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    photo_datas = child.get("photo_data", [])
                    table = doc.add_table(rows=1, cols=1, style="Table Grid")
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell = table.cell(0, 0)
                    cell.width = Cm(18.9)
                    paragraph = cell.paragraphs[0]
                    for photo_data in photo_datas:
                        run = paragraph.add_run()
                        try:
                            run.add_picture(current_app.config['UPLOAD_FOLDER']+"/"+photo_data["filepath"].split("solar_static/")[-1], height=Inches(1.75))
                            # run.add_picture("ntust.png", height=Inches(1.75))
                        except:
                            print("no picture, path："+current_app.config['UPLOAD_FOLDER']+"/"+photo_data["filepath"].split("solar_static/")[-1])
                            run.text = "資料庫無相片"
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = paragraph.add_run("  ")
            if index<len(photo_dict_list)-1:
                doc.add_page_break()
            
    elif list_name in ["alarm_list"]:
        for target in target_list:
            for user_select in target.get("cause_data", {}).get("user_select", []):
                if len(user_select.get("photo_data", [])) > 0:
                    title_name = target.get("event", "")+"-"+user_select.get("group")+"/"+user_select.get("event")
                    if title_name not in title_list:
                        title_list.append(title_name)
                    photo_dict[title_name] = user_select.get("photo_data", [])
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run("異常排除")
        run.font.size = Pt(20)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        for index, title in enumerate(title_list):
            table = doc.add_table(rows=1, cols=1, style="Table Grid")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            paragraph = table.cell(0, 0).paragraphs[0]
            run = paragraph.add_run(title)
            run.font.size = Pt(14)
            table.cell(0, 0).width = Cm(18.9)
            table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table = doc.add_table(rows=1, cols=1, style="Table Grid")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for photo in photo_dict.get(title, []):
                cell = table.cell(0, 0)
                cell.width = Cm(18.9)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                try:
                    run.add_picture(current_app.config['UPLOAD_FOLDER']+"/"+photo["filepath"].split("solar_static/")[-1], height=Inches(1.75))
                    # run.add_picture("ntust.png", height=Inches(1.75))
                except:
                    print("no picture, path："+current_app.config['UPLOAD_FOLDER']+"/"+photo["filepath"].split("solar_static/")[-1])
                    run.text = "資料庫無相片"
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run("  ")

def title_page(doc, title_dict, logo_picture, info_picture):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    try:
        run.add_picture(logo_picture, height=Inches(2.0))
    except:
        print("no logo, path：{}".format(logo_picture))
    paragraph = doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("維 運 報 告 ")
    run.font.underline = True
    run.font.size = Pt(48)
    run = paragraph.add_run("O&M REPORT")
    run.font.underline = True
    run.font.size = Pt(36)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(title_dict["zh-tw"])
    run.font.size = Pt(28)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(title_dict["en-us"])
    run.font.size = Pt(24)

    for _ in range(6):
        paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(info_picture, height=Inches(3.0))
    except:
        print("no info_picture, path：{}".format(info_picture))
    doc.add_page_break()

def report_dispatch_maohong(db, request_dict):
    try:
        dispatch_data = get_dispatch_data(db, request_dict)
        # print(dispatch_data)

        stage = {
            "wait_for_priority": "等待優先度排序",
            "wait_for_take": "待接單",
            "merged": "已合併",
            "took_wait_date_enter": "等待輸入派工日期",
            "wait_for_dispatch": "等待派工",
            "wait_admin_confirm_date": "等待管理人員確認派工日期",
            "dispatched_wait_for_review": "等待AI驗收",
            "auto_reviewed_wait_for_manual": "等待管理人員協助驗收",
            "review_failed": "驗收失敗",
            "dispatch_finish": "已完成工單"
        }

        last_time, entry_time = dispatch_last_dispatch_time(db, dispatch_data)

        start_date = dispatch_data.get("station_data", {}).get("plant_data", {}).get("start_date", "")
        start_date = start_date.replace("-", "/")
        module_name, inverter_name = dispatch_find_module_inverter(db, dispatch_data)
        maohong_station_info = {
            "dispatch_number": dispatch_data.get("name", ""),
            "plant_name": dispatch_data.get("station_data", {}).get("plant_data", {}).get("pv_plant_client_code", "")+"/"+dispatch_data.get("station_data", {}).get("plant_data", {}).get("name", ""),
            "entry_time": entry_time,
            "worker": dispatch_data.get("maintainer_data", {}).get("name", ""),
            "capacity": str(dispatch_data.get("station_data", {}).get("plant_data", {}).get("capacity", "---")),
            "model": module_name,
            "inverter": inverter_name,
            "angle": dispatch_data.get("station_data", {}).get("plant_data", {}).get("angle", "---"),
            "start_date": start_date,
            "last_date": last_time
        }

        maintain_data = get_maintaince(db, dispatch_data)
        maintain_info = {
            "regular": maintain_data[0],
            "wash": maintain_data[1],
            "alarm": maintain_data[2]
        }

        title_info = {
            "zh-tw": maintain_data[3],
            "en-us": maintain_data[4]
        }

        doc = Document()
        #初始設定
        margins = {"left": 1.38, "right": 1.69, "top": 1.13, "bottom": 1.08}
        document_initial(doc, margins)


        # dispatch_maohong_header(doc, "maohong_logo.png")
        dispatch_maohong_header(doc, current_app.config['UPLOAD_FOLDER']+"/images/maohong_logo.jpg")

        # title_page(doc, title_info, "maohong_logo_picture.jpg", "maohong_info_picture.jpg")
        title_page(doc, title_info, current_app.config['UPLOAD_FOLDER']+"/images/maohong_logo_picture.jpg", current_app.config['UPLOAD_FOLDER']+"/images/maohong_info_picture.jpg")

        dispatch_maohong_record_information(doc, maohong_station_info, maintain_info)

        doc.add_page_break()

        for type_name in dispatch_data.get("type", []):
            target_list = dispatch_data.get("{}_list".format(type_name), [])
            list_name = "{}_list".format(type_name)
            list_table(doc, target_list, list_name)

        for index, type_name in enumerate(dispatch_data.get("type", [])):
            target_list = dispatch_data.get("{}_list".format(type_name), [])
            list_name = "{}_list".format(type_name)
            list_picture(doc, target_list, list_name)
            if len(dispatch_data.get("type", []))>1 and index<len(dispatch_data.get("type", []))-1:
                doc.add_page_break()

        filename = "派工報表-{}.docx".format(dispatch_data.get("name", ""))
        print(filename)
        print("dispatch report finish")
        # doc.save("報表.docx")
        # doc.save(filename)
        doc_name = quote(filename)

        # Create in-memory buffer
        file_stream = BytesIO()
        # Save the .docx to the buffer
        doc.save(file_stream)
        # Reset the buffer's file-pointer to the beginning of the file
        file_stream.seek(0)
        rv = send_file(
            file_stream,
            as_attachment=True,
            attachment_filename=doc_name,
            download_name=doc_name
        )
        return rv
        # doc.save(filename)
    except Exception as e:
        print("docx_produce_maohong:")
        print(exception_detail(e))
#-------------------------------------------------------------------------------------------
    